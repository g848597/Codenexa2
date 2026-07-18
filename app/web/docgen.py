"""Рендер готового документа (текст с уже подставленными значениями) в
PDF и DOCX. Один источник разметки для обоих форматов: заголовок по
центру, основной текст абзацами, при наличии профиля пользователя —
блок с датой/местом и подписью (плюс картинка подписи/логотип, если
загружены — см. /api/profile/logo, /api/profile/signature в docs.py).

Сознательно не тянем сюда HTML/CSS-рендер (WeasyPrint и т.п.) — документы
это простой деловой текст, reportlab/python-docx дают полный контроль над
вёрсткой без системных зависимостей (WeasyPrint требует Pango/Cairo на
хосте, что лишняя головная боль на Railway/VPS).
"""
import io
import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

_INK = colors.HexColor("#1c2230")
_MUTED = colors.HexColor("#6b7280")

_INK_DOCX = RGBColor(0x1C, 0x22, 0x30)
_MUTED_DOCX = RGBColor(0x6B, 0x72, 0x80)


def _paragraphs(body_text: str) -> list[str]:
    """Разбивает текст документа на абзацы по пустой строке; одиночные
    переносы внутри абзаца сохраняются как <br/> (PDF) / мягкий перенос (DOCX)."""
    raw = (body_text or "").replace("\r\n", "\n").strip("\n")
    parts = [p.strip("\n") for p in raw.split("\n\n")]
    return [p for p in parts if p.strip()]


def _safe_image_path(path: str | None) -> str | None:
    if not path:
        return None
    if not os.path.isfile(path):
        return None
    return path


# --------------------------------------------------------------------------
# PDF (reportlab)
# --------------------------------------------------------------------------

def build_pdf(title: str, body_text: str, profile: dict | None = None) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.4 * cm, rightMargin=2.4 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=title,
    )

    title_style = ParagraphStyle(
        "DocTitle", fontName="Helvetica-Bold", fontSize=14.5, leading=18,
        alignment=TA_CENTER, textColor=_INK, spaceAfter=18,
    )
    body_style = ParagraphStyle(
        "DocBody", fontName="Helvetica", fontSize=10.8, leading=16.5,
        alignment=TA_JUSTIFY, textColor=_INK, spaceAfter=10,
    )
    muted_style = ParagraphStyle(
        "Muted", fontName="Helvetica", fontSize=9.5, leading=13,
        alignment=TA_CENTER, textColor=_MUTED,
    )

    story = []

    logo_path = _safe_image_path((profile or {}).get("logo_path_abs"))
    if logo_path:
        try:
            img = RLImage(logo_path, width=3.2 * cm, height=3.2 * cm, kind="proportional")
            story.append(img)
            story.append(Spacer(1, 8))
        except Exception:
            pass

    story.append(Paragraph(_escape(title.upper()), title_style))

    for para in _paragraphs(body_text):
        html_para = _escape(para).replace("\n", "<br/>")
        story.append(Paragraph(html_para, body_style))

    story.append(Spacer(1, 26))
    story.extend(_signature_block_pdf(profile))

    story.append(Spacer(1, 18))
    story.append(Paragraph(
        f"Документ сформирован в CodeNexa AI Docs · {datetime.now().strftime('%d.%m.%Y')}",
        muted_style,
    ))

    doc.build(story)
    return buf.getvalue()


def _signature_block_pdf(profile: dict | None):
    profile = profile or {}
    now = datetime.now()
    city = profile.get("city") or "_______________"
    name = profile.get("signature_name") or profile.get("full_name") or ""

    sign_style = ParagraphStyle(
        "SignLabel", fontName="Helvetica", fontSize=10, leading=14, textColor=_INK,
    )

    left = Paragraph(
        f"г. {_escape(city)}&nbsp;&nbsp;&nbsp;«{now.strftime('%d')}» {_month_ru(now)} {now.strftime('%Y')} г.",
        sign_style,
    )

    sig_path = _safe_image_path(profile.get("signature_path_abs"))
    if sig_path:
        try:
            sig_cell = RLImage(sig_path, width=3.6 * cm, height=1.4 * cm, kind="proportional")
        except Exception:
            sig_cell = Paragraph("Подпись: _______________", sign_style)
    else:
        sig_cell = Paragraph("Подпись: _______________", sign_style)

    table = Table([[left, sig_cell]], colWidths=[9 * cm, 7 * cm])
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
    ]))

    rows = [table]
    if name:
        rows.append(Spacer(1, 4))
        rows.append(Paragraph(_escape(name), ParagraphStyle(
            "SignName", fontName="Helvetica", fontSize=9.5, leading=12,
            alignment=2, textColor=_MUTED,  # 2 = TA_RIGHT
        )))
    return rows


def _month_ru(dt: datetime) -> str:
    months = ["января", "февраля", "марта", "апреля", "мая", "июня",
              "июля", "августа", "сентября", "октября", "ноября", "декабря"]
    return months[dt.month - 1]


def _escape(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                .replace("&lt;br/&gt;", "<br/>"))


# --------------------------------------------------------------------------
# DOCX (python-docx)
# --------------------------------------------------------------------------

def build_docx(title: str, body_text: str, profile: dict | None = None) -> bytes:
    profile = profile or {}
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = _INK_DOCX

    logo_path = _safe_image_path(profile.get("logo_path_abs"))
    if logo_path:
        try:
            doc.add_picture(logo_path, width=Cm(3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(15)
    h.paragraph_format.space_after = Pt(18)

    for para_text in _paragraphs(body_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.3
        lines = para_text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)

    doc.add_paragraph()

    today = datetime.now()
    city = profile.get("city") or "_______________"
    date_line = doc.add_paragraph()
    date_line.add_run(f"г. {city}    «{today.strftime('%d')}» {_month_ru(today)} {today.strftime('%Y')} г.")

    sig_path = _safe_image_path(profile.get("signature_path_abs"))
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if sig_path:
        try:
            doc.add_picture(sig_path, width=Cm(3.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            sig_para.add_run("Подпись: _______________")
    else:
        sig_para.add_run("Подпись: _______________")

    name = profile.get("signature_name") or profile.get("full_name")
    if name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = name_para.add_run(f"({name})")
        run.font.size = Pt(10)
        run.font.color.rgb = _MUTED_DOCX

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"Документ сформирован в CodeNexa AI Docs · {today.strftime('%d.%m.%Y')}")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = _MUTED_DOCX

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
